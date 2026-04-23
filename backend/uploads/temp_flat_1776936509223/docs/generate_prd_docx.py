from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import os

doc = Document()

def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading_custom(text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt([18, 14, 12][min(level-1, 2)])
    run.bold = True
    return p

def add_para_custom(text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def parse_inline_formatting(text):
    parts = []
    pattern = r'\*\*(.*?)\*\*|`(.*?)`|\*(.*?)\*'
    pos = 0
    for m in re.finditer(pattern, text):
        if m.start() > pos:
            parts.append(('normal', text[pos:m.start()]))
        if m.group(1):
            parts.append(('bold', m.group(1)))
        elif m.group(2):
            parts.append(('code', m.group(2)))
        elif m.group(3):
            parts.append(('italic', m.group(3)))
        pos = m.end()
    if pos < len(text):
        parts.append(('normal', text[pos:]))
    return parts

def add_formatted_para(text):
    p = doc.add_paragraph()
    parts = parse_inline_formatting(text)
    for style, content in parts:
        run = p.add_run(content)
        set_run_font(run)
        if style == 'bold':
            run.bold = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.color.rgb = RGBColor(0xd7, 0x3a, 0x49)
    return p

with open('docs/PRD.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_buffer = []
in_table = False
table_lines = []

for line in lines:
    line = line.rstrip('\n')
    if line.startswith('```'):
        if in_code:
            if code_buffer:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                run.font.size = Pt(10)
                code_buffer = []
            in_code = False
        else:
            in_code = True
        continue
    if in_code:
        code_buffer.append(line)
        continue

    if line.startswith('|'):
        table_lines.append(line)
        in_table = True
        continue
    else:
        if in_table:
            if len(table_lines) >= 2:
                rows_data = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip('|').split('|')]
                    rows_data.append(cells)
                if rows_data and set(rows_data[1]) <= {'', '-', ':', ' '}:
                    rows_data.pop(1)
                if rows_data:
                    table = doc.add_table(rows=1, cols=len(rows_data[0]))
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    for i, cell_text in enumerate(rows_data[0]):
                        hdr[i].text = cell_text
                    for row_data in rows_data[1:]:
                        row = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            if i < len(row):
                                row[i].text = cell_text
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                                    run.font.size = Pt(10)
            table_lines = []
            in_table = False

    if line.startswith('# '):
        add_heading_custom(line[2:], level=1)
    elif line.startswith('## '):
        add_heading_custom(line[3:], level=2)
    elif line.startswith('### '):
        add_heading_custom(line[4:], level=3)
    elif line.startswith('#### '):
        add_heading_custom(line[5:], level=4)
    elif line.startswith('- [ ] '):
        add_formatted_para('□ ' + line[6:])
    elif line.startswith('- '):
        add_formatted_para('• ' + line[2:])
    elif line.startswith('> '):
        p = add_formatted_para(line[2:])
        p.paragraph_format.left_indent = Inches(0.3)
    elif line.strip() == '':
        doc.add_paragraph()
    elif line.startswith('**PRD 结束**'):
        add_para_custom(line.strip('*'), bold=True)
    else:
        add_formatted_para(line)

doc.save('docs/PRD.docx')
print('docs/PRD.docx generated successfully')
