import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def parse_bold_italic(run, text):
    """简单解析 **bold** 和 *italic*，添加到 run"""
    # 由于 python-docx 的 add_run 在段落级别更容易处理，
    # 这里我们直接返回纯文本，样式解析比较复杂，先保持文本原样可读
    run.text = text

def add_paragraph(doc, text, style='Normal', bold=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    if bold:
        r.bold = True
    return p

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # 设置默认中文字体（Windows 上 SimSun 或 Microsoft YaHei）
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        i += 1

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph(style='Normal')
                r = p.add_run('\n'.join(code_lines))
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # 灰色背景通过底纹设置较复杂，这里用边框简单区分
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # 空行
        if not line.strip():
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(18)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            elif level == 2:
                p = doc.add_heading(text, level=1)
                for r in p.runs:
                    r.font.size = Pt(14)
                    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            elif level == 3:
                p = doc.add_heading(text, level=2)
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            else:
                p = doc.add_heading(text, level=3)
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # 分隔线
        if line.strip() == '---' or line.strip() == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0)
            r = p.add_run('_' * 60)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            continue

        # 表格开始
        if line.strip().startswith('|') and i < len(lines) and lines[i].strip().startswith('|'):
            # 收集表格所有行
            table_lines = [line]
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1

            # 过滤掉 Markdown 表格分隔行（只包含 - 和 | 和 : 和空格）
            data_lines = []
            for tl in table_lines:
                content = tl.strip()
                # 去掉首尾的 |
                if content.startswith('|'): content = content[1:]
                if content.endswith('|'): content = content[:-1]
                cells = [c.strip() for c in content.split('|')]
                # 判断是否分隔行
                if all(re.match(r'^[-:]+$', c.replace(' ', '')) for c in cells if c):
                    continue
                data_lines.append(cells)

            if data_lines:
                table = doc.add_table(rows=len(data_lines), cols=len(data_lines[0]))
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(data_lines):
                    row = table.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx >= len(row.cells):
                            break
                        cell = row.cells[c_idx]
                        cell.text = cell_text
                        # 首行加粗
                        if r_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                        else:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
            continue

        # 列表项
        list_match = re.match(r'^(\s*)([-*])\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            r = p.add_run(text)
            r.font.size = Pt(10.5)
            continue

        # 数字列表
        num_list_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if num_list_match:
            indent = len(num_list_match.group(1))
            text = num_list_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.1)
            r = p.add_run(text)
            r.font.size = Pt(10.5)
            continue

        # 引用块
        quote_match = re.match(r'^>\s*(.*)$', line)
        if quote_match:
            text = quote_match.group(1).strip()
            p = doc.add_paragraph(style='Intense Quote')
            r = p.add_run(text)
            r.font.size = Pt(10.5)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # 普通段落
        p = doc.add_paragraph(style='Normal')
        r = p.add_run(line.strip())
        r.font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"Saved to {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx('docs/PRD.md', 'docs/PRD.docx')
