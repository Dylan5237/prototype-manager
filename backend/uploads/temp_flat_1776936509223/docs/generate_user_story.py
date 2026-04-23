from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os

doc = Document()

def set_run_font(run, size=11, bold=False):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold

def add_heading_custom(text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size=Pt([16, 14, 12][min(level-1, 2)]), bold=True)
    return p

def add_para_custom(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p

# ======================
# 用户故事文档内容
# ======================
add_heading_custom('用户故事：业务域 CI/CD 协作流程', level=1)
add_para_custom('')

add_heading_custom('1. 背景与目标', level=2)
add_para_custom('在医院信息系统的业务域持续演进过程中，每个业务域（如医嘱、草药处方、皮试记录等）都需要独立的版本管理和 CI/CD 流水线。本用户故事描述了“业务域版本管理”页面如何支撑这一完整协作流程。')

add_heading_custom('2. 用户故事列表', level=2)

stories = [
    ('US-01 业务域版本概览', '作为业务域开发者，我希望在“概览”页看到当前版本的依赖范围、关联需求、变更记录和测试状态，以便快速了解版本健康度。', [
        '打开“业务域版本管理”页面，默认进入“概览”标签页。',
        '系统展示当前版本号、分支、依赖范围、关联需求列表。',
        '每项需求显示进度状态（待开发 / 测试中 / 已完成）。',
        '可切换“最新 / 历史”视图查看过往版本。'
    ]),
    ('US-02 提交测试（创建发布）', '作为业务域开发者，当我完成一组需求后，我希望点击“提交测试”生成一个测试版本，以便触发 CI 构建。', [
        '在“概览”或“发布记录”页点击“提交测试”。',
        '弹出对话框，选择分支（main / bugfix）。',
        '系统自动生成语义化版本号（如 1.2.3）。',
        '填写发布说明，勾选本次关联的需求。',
        '点击“提交测试”后，系统记录发布条目，状态为“待测试”。',
        '后台触发 CI 流水线：编译 → 打包 → 生成版本产物。'
    ]),
    ('US-03 添加/更新依赖', '作为业务域开发者，我需要声明当前版本依赖的其他业务域版本范围，以避免运行时兼容性问题。', [
        '在“概览”页点击“添加依赖”。',
        '选择被依赖的业务域。',
        '分别选择“最小依赖版本”和“最大依赖版本”。',
        '系统校验：若当前环境版本不在 [min, max] 范围内，标记状态为“不兼容”。',
        '保存后，依赖信息同步写入版本元数据，供下游 CI 解析。'
    ]),
    ('US-04 QA 验证与自动测试', '作为测试人员，我希望在测试环境部署后查看自动化测试结果，并对失败项提交 BUG。', [
        'CI 构建完成后，自动部署到测试环境。',
        'QA 在“发布记录”页看到版本状态变为“测试中”。',
        '点击“查看详情”可查看构建日志、单元测试报告、接口覆盖率。',
        '若发现缺陷，QA 可直接点击“新建 BUG”关联到当前版本。',
        '开发者在“BUG修复”页看到待处理列表。'
    ]),
    ('US-05 版本发布与下游消费', '作为下游业务域开发者，我希望及时获取依赖域的最新可用版本，并在兼容范围内自动升级。', [
        '当前版本通过 QA 验证后，状态变为“已通过”。',
        '运维/负责人点击“发布到仓库”，状态变为“已发布”。',
        '下游域的 CI 在构建时读取依赖版本范围。',
        '若仓库中存在满足 [min, max] 的新版本，自动拉取并编译。',
        '若新版本超出 max，则保持使用当前版本并提示升级。'
    ]),
    ('US-06 反馈修复循环', '作为开发者，当测试或下游反馈问题时，我希望快速定位并发布补丁版本。', [
        '在“BUG修复”页查看所有未关闭的 BUG。',
        '选中一个 BUG，点击“修复”切换分支到 bugfix。',
        '提交代码后，再次点击“提交测试”生成补丁版本（patch +1）。',
        'QA 对补丁版本进行回归验证。',
        '验证通过后合并回 main，并发布。',
        '下游域自动感知兼容范围内的补丁更新。'
    ])
]

for title, desc, steps in stories:
    add_heading_custom(title, level=3)
    add_para_custom(f'描述：{desc}')
    add_para_custom('验收标准：', bold=True)
    for step in steps:
        add_para_custom(f'    • {step}')
    add_para_custom('')

add_heading_custom('3. 流程图说明', level=2)
add_para_custom('下图展示了从“建模开发 → 提交测试 → CI构建 → QA验证 → 发布 → 下游消费 → 反馈修复”的完整闭环。')
if os.path.exists('docs/cicd_flow.png'):
    doc.add_picture('docs/cicd_flow.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('docs/用户故事-业务域CI-CD协作流程.docx')
print('docs/用户故事-业务域CI-CD协作流程.docx generated successfully')

# ======================
# 绘制 CI/CD 流程图
# ======================
width, height = 800, 1200
img = Image.new('RGB', (width, height), (255, 255, 255))
draw = ImageDraw.Draw(img)

try:
    font = ImageFont.truetype("msyh.ttc", 18)
    font_small = ImageFont.truetype("msyh.ttc", 14)
except Exception:
    font = ImageFont.load_default()
    font_small = ImageFont.load_default()

def draw_rounded_rect(draw, xy, fill, outline=None, radius=10, width=2):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)

def draw_arrow(draw, x1, y1, x2, y2, color=(100, 100, 100), width=2):
    draw.line((x1, y1, x2, y2), fill=color, width=width)
    # simple arrowhead
    if y2 > y1:
        draw.polygon([(x2, y2), (x2-5, y2-8), (x2+5, y2-8)], fill=color)
    elif y2 < y1:
        draw.polygon([(x2, y2), (x2-5, y2+8), (x2+5, y2+8)], fill=color)
    elif x2 > x1:
        draw.polygon([(x2, y2), (x2-8, y2-5), (x2-8, y2+5)], fill=color)
    else:
        draw.polygon([(x2, y2), (x2+8, y2-5), (x2+8, y2+5)], fill=color)

def get_text_size(draw, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]

nodes = [
    "1. 业务域建模开发",
    "2. 提交测试 / 创建发布",
    "3. CI 构建 & 版本生成",
    "4. 测试环境自动部署",
    "5. QA 验证 / 自动测试",
    "6. 合并到主干",
    "7. 发布到仓库",
    "8. 下游域依赖拉取",
    "9. 反馈修复 & 补丁发布"
]

box_w, box_h = 280, 50
start_y = 60
gap_y = 40
center_x = width // 2

positions = []
for i, text in enumerate(nodes):
    x = center_x - box_w // 2
    y = start_y + i * (box_h + gap_y)
    positions.append((x, y, x + box_w, y + box_h))
    # shadow
    draw.rounded_rectangle([x+3, y+3, x+box_w+3, y+box_h+3], radius=12, fill=(200, 200, 200))
    # box
    fill_color = (59, 130, 246) if i in [0, 4, 8] else (255, 255, 255)
    outline_color = (59, 130, 246)
    draw_rounded_rect(draw, [x, y, x+box_w, y+box_h], fill=fill_color, outline=outline_color, radius=12, width=2)
    # text
    text_color = (255, 255, 255) if i in [0, 4, 8] else (33, 33, 33)
    tw, th = get_text_size(draw, text, font)
    tx = center_x - tw // 2
    ty = y + (box_h - th) // 2
    draw.text((tx, ty), text, fill=text_color, font=font)

# draw arrows between main nodes
for i in range(len(positions) - 1):
    x1 = center_x
    y1 = positions[i][3]
    x2 = center_x
    y2 = positions[i+1][1]
    draw_arrow(draw, x1, y1, x2, y2)

# side annotations
annotations = [
    ("需求建模", 0, -1),
    ("语义版本生成", 1, 1),
    ("依赖范围声明", 2, -1),
    ("自动测试报告", 4, 1),
    ("兼容性校验", 7, -1),
    ("补丁版本 +1", 8, 1),
]

for text, idx, side in annotations:
    tw, th = get_text_size(draw, text, font_small)
    if side == -1:
        ax = positions[idx][0] - tw - 20
    else:
        ax = positions[idx][2] + 20
    ay = positions[idx][1] + (box_h - th) // 2
    draw.text((ax, ay), text, fill=(107, 114, 128), font=font_small)
    # dotted line to box
    bx = positions[idx][0] if side == -1 else positions[idx][2]
    by = positions[idx][1] + box_h // 2
    ex = ax + tw if side == -1 else ax
    ey = ay + th // 2
    for k in range(0, abs(ex - bx), 6):
        sx = bx + k * side
        draw.line((sx, by, sx + 3 * side, by), fill=(150, 150, 150), width=1)

# feedback loop arrow from last to second
# draw a curved arrow from bottom of last box back to right of second box
x1 = positions[-1][2]
y1 = positions[-1][1] + box_h // 2
x2 = positions[1][2] + 60
y2 = positions[1][1] + box_h // 2
draw.line((x1, y1, x2, y1), fill=(217, 48, 37), width=2)
draw.line((x2, y1, x2, y2), fill=(217, 48, 37), width=2)
draw.polygon([(x2, y2), (x2-5, y2+8), (x2+5, y2+8)], fill=(217, 48, 37))
loop_text = "反馈闭环"
tw, th = get_text_size(draw, loop_text, font_small)
draw.text((x2 + 8, y2 - th // 2), loop_text, fill=(217, 48, 37), font=font_small)

# title
title = "业务域 CI/CD 协作流程图"
tw, th = get_text_size(draw, title, font)
draw.text((center_x - tw // 2, 10), title, fill=(33, 33, 33), font=font)

img.save('docs/cicd_flow.png')
print('docs/cicd_flow.png generated successfully')
